import argparse
from docx import Document
import os
from datetime import datetime

def extract_images(docx_path, output_folder='extracted_images', backup=True, remove_images=False):
    # 确保输入文件存在
    if not os.path.exists(docx_path):
        print(f"错误: 找不到文件 '{docx_path}'")
        return
    
    # 创建备份
    if backup:
        backup_path = f"{docx_path}.backup"
        try:
            import shutil
            shutil.copy2(docx_path, backup_path)
            print(f"已创建备份文件: {backup_path}")
        except Exception as e:
            print(f"警告: 创建备份失败: {str(e)}")
            return
    
    # 创建输出文件夹
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
        print(f"创建输出文件夹: {output_folder}")
    
    # 打开word文档
    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"错误: 无法打开文件 '{docx_path}': {str(e)}")
        return
    
    # 用于记录提取的图片数量
    image_count = 0
    
    # 提取并删除图片
    rels_to_delete = []
    for rel in list(doc.part.rels.values()):
        # 检查是否是图片
        if "image" in rel.target_ref:
            try:
                # 获取图片数据
                image_data = rel.target_part.blob
                
                # 生成图片文件名
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                image_filename = f'image_{timestamp}_{image_count}.png'
                image_path = os.path.join(output_folder, image_filename)
                
                # 保存图片
                with open(image_path, 'wb') as f:
                    f.write(image_data)
                image_count += 1
                print(f'已保存图片: {image_filename}')
                
                # 只在需要删除图片时才记录关系ID
                if remove_images:
                    rels_to_delete.append(rel.rId)
                
            except Exception as e:
                print(f"警告: 处理图片时出错: {str(e)}")
    
    # 只在需要删除图片时执行删除操作
    if remove_images:
        # 删除图片关系
        for rId in rels_to_delete:
            del doc.part.rels[rId]
        
        # 完全清理文档中的图片元素
        for paragraph in doc.paragraphs:
            p = paragraph._element
            
            # 使用 findall 替代 xpath
            ns = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                'v': 'urn:schemas-microsoft-com:vml',
                'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
                'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'
            }
            
            # 删除 drawing 元素
            for element in p.findall('.//w:drawing', ns):
                parent = element.getparent()
                if parent is not None:
                    parent.remove(element)
            
            # 删除 pict 元素
            for element in p.findall('.//w:pict', ns):
                parent = element.getparent()
                if parent is not None:
                    parent.remove(element)
            
            # 删除 imagedata 元素
            for element in p.findall('.//v:imagedata', ns):
                parent = element.getparent()
                if parent is not None:
                    parent.remove(element)
                    
            # 删除 shape 元素
            for element in p.findall('.//v:shape', ns):
                parent = element.getparent()
                if parent is not None:
                    parent.remove(element)
            
            # 清理空的运行元素
            for run in p.findall('.//w:r', ns):
                if len(run) == 0 or (len(run) == 1 and run[0].tag.endswith('}rPr')):
                    parent = run.getparent()
                    if parent is not None:
                        parent.remove(run)
        
        # 保存修改后的文档
        try:
            doc.save(docx_path)
            print(f"\n文档已更新: {docx_path}")
        except Exception as e:
            print(f"错误: 保存文档失败: {str(e)}")
            return
        
        print(f'总共提取了 {image_count} 张图片')
        print(f'图片保存在文件夹: {output_folder}')
    else:
        print("\n文档保持不变，仅提取了图片")


def extract_text(docx_path, save=False, save_path='extracted_text'):
    if not os.path.exists(docx_path):
        print(f"错误: 找不到文件 '{docx_path}'")
        return None
    
    try:
        doc = Document(docx_path)
        markdown_text = []
        
        # 找到第一个空白段落的位置（用于区分标题和正文）
        blank_para_index = -1
        for i, para in enumerate(doc.paragraphs):
            if not para.text.strip() and i > 0 and doc.paragraphs[i-1].text.strip():
                blank_para_index = i
                break
        
        # 处理段落
        for i in range(len(doc.paragraphs)):
            paragraph = doc.paragraphs[i]
            text = paragraph.text.strip()
            if not text:  # 跳过空段落
                continue
            
            # 标题部分处理
            if blank_para_index > 0 and i < blank_para_index:
                if i == 0:  # 第一段为主标题
                    markdown_text.append(f'# {text}\n\n')
                else:
                    markdown_text.append(f'# {text}\n\n')
            # 正文部分处理
            else:
                # 检查是否以中文数字加顿号开头（二级标题）
                if any(text.startswith(f"{num}、") for num in ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]):
                    markdown_text.append(f'## {text}\n\n')
                
                # 检查是否包含（一）到（十）（三级标题）
                elif any(f"（{num}）" in text for num in ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]):
                    markdown_text.append(f'### {text}\n\n')
                
                # 检查是否以阿拉伯数字加点开头（三级标题）
                elif any(text.startswith(f"{str(num)}.") for num in range(1, 11)):
                    markdown_text.append(f'### {text}\n\n')
                
                # 检查是否包含"一是"到"十是"（加粗处理）
                elif any(f"{num}是" in text for num in ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]):
                    # 找到"X是"的位置和句号的位置
                    for num in ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]:
                        pattern = f"{num}是"
                        if pattern in text:
                            start_idx = text.find(pattern)
                            end_idx = text.find("。", start_idx)
                            if end_idx != -1:
                                # 将"X是"到句号的部分加粗
                                bold_part = text[start_idx:end_idx + 1]
                                markdown_text.append(f'{text[:start_idx]}**{bold_part}**{text[end_idx + 1:]}\n\n')
                            else:
                                markdown_text.append(f'{text}\n\n')
                            break
                else:
                    markdown_text.append(f'{text}\n\n')
        
        # 合并所有文本
        final_text = ''.join(markdown_text)
        
        # 如果需要保存文件
        if save:
            if not os.path.exists(save_path):
                os.makedirs(save_path)
            
            base_name = os.path.splitext(os.path.basename(docx_path))[0]
            output_file = os.path.join(save_path, f'{base_name}.md')
            
            try:
                with open(output_file, 'w', encoding='utf-8') as f:
                    f.write(final_text)
                print(f'文本已保存到: {output_file}')
            except Exception as e:
                print(f"错误: 保存文件失败: {str(e)}")
        
        return final_text
        
    except Exception as e:
        print(f"错误: 处理文档时出错: {str(e)}")
        return None

def test_extract_text():
    extract_text('test/test.docx', save=True, save_path='test')
    
if __name__ == '__main__':
    test_extract_text()