import argparse
import os
from pathlib import Path
import yaml
from docx import Document
from docx.shared import Pt, Length
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def load_rules(rules_path):
    """加载自定义规则文件"""
    try:
        with open(rules_path, 'r', encoding='utf-8') as f:
            rules = yaml.safe_load(f)
        return rules
    except Exception as e:
        print(f"Error loading rules file: {e}")
        return None

def format_title_section(doc_path, custom_rules=None):
    # 使用默认规则或自定义规则
    rules = custom_rules if custom_rules else {
        'title': {
            'font': '宋体',
            'size': 22,
            'bold': True,
            'align': 'center'
        },
        'content': {
            'default': {
                'font': '仿宋',
                'size': 16,
                'bold': False,
                'align': 'justify'
            },
            'number_dot': {  # 一、二、等开头
                'font': '黑体',
                'bold': False
            },
            'number_brackets': {  # （一）（二）等
                'font': '楷体',
                'bold': True
            },
            'number_is': {  # 一是 二是等
                'font': '楷体',
                'bold': True
            },
            'arabic_number': {  # 1. 2. 等开头
                'font': '楷体',
                'bold': True
            }
        }
    }
    
    # 打开Word文档
    doc = Document(doc_path)
    
    # 先清除所有段落的格式
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():  # 如果段落有文本
            # 保存文本并重新添加，清除所有格式
            text = paragraph.text.strip()
            paragraph.clear()  # 清除所有runs
            run = paragraph.add_run(text)  # 重新添加文本
            # 清除段落格式
            paragraph.paragraph_format.first_line_indent = Pt(0)
            paragraph.paragraph_format.left_indent = Pt(0)
            paragraph.paragraph_format.right_indent = Pt(0)
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
    
    # 找到第一个空白段落的位置
    blank_para_index = -1
    for i, para in enumerate(doc.paragraphs):
        if not para.text.strip():  # 检查是否为空白段落
            # 确保空白段落上方有文字内容
            if i > 0 and doc.paragraphs[i-1].text.strip():
                blank_para_index = i
                break
    
    # 如果找到符合条件的空白段落，处理所有段落
    if blank_para_index > 0:
        # 处理标题部分（空白段落之前）
        for i in range(blank_para_index):
            paragraph = doc.paragraphs[i]
            if paragraph.text.strip():  # 只处理非空段落
                original_text = paragraph.text
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                # 设置标题行距
                paragraph.paragraph_format.line_spacing = Pt(28.5)
                paragraph.clear()
                run = paragraph.add_run(original_text)
                run.font.name = rules['title']['font'] 
                run._element.rPr.rFonts.set(qn('w:eastAsia'), rules['title']['font'])
                run.font.size = Pt(rules['title']['size'])
                run.font.bold = rules['title']['bold']
        
        # 处理正文部分（空白段落之后）
        # 创建新的段落列表，只包含非空段落
        content_paragraphs = []
        for i in range(blank_para_index + 1, len(doc.paragraphs)):
            paragraph = doc.paragraphs[i]
            if paragraph.text.strip():  # 只保留非空段落
                content_paragraphs.append(paragraph)
        
        # 删除原文档中从空白段落后的所有段落
        for _ in range(len(doc.paragraphs) - (blank_para_index + 1)):
            p = doc.paragraphs[blank_para_index + 1]._element
            p.getparent().remove(p)
        
        # 重新添加格式化后的非空段落
        for paragraph in content_paragraphs:
            original_text = paragraph.text
            new_para = doc.add_paragraph()
            new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # 两端对齐
            new_para.paragraph_format.line_spacing = Pt(28.5)

            def apply_format(run, style_key):
                """应用指定的格式"""
                run.font.name = rules['content'][style_key]['font']
                run._element.rPr.rFonts.set(qn('w:eastAsia'), rules['content'][style_key]['font'])
                run.font.bold = rules['content'][style_key]['bold']
                run.font.size = Pt(rules['content']['default']['size'])

            # 检查是否包含（一）到（十）的句子
            if any(f"（{num}）" in original_text for num in ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]):
                for num in ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]:
                    pattern = f"（{num}）"
                    if pattern in original_text:
                        start_idx = original_text.find(pattern)
                        end_idx = original_text.find("。", start_idx)
                        if end_idx != -1:
                            if start_idx > 0:
                                run = new_para.add_run("    " + original_text[:start_idx])
                                apply_format(run, 'default')
                            run = new_para.add_run(original_text[start_idx:end_idx + 1])
                            apply_format(run, 'number_brackets')
                            if end_idx + 1 < len(original_text):
                                run = new_para.add_run(original_text[end_idx + 1:])
                                apply_format(run, 'default')
                            else:
                                run = new_para.add_run("      " + original_text)
                                apply_format(run, 'default')
                            break
            # 检查是否以阿拉伯数字加点开头
            elif any(original_text.startswith(f"{str(num)}.") for num in range(1, 11)):
                for num in range(1, 11):
                    pattern = f"{str(num)}."
                    if original_text.startswith(pattern):
                        end_idx = original_text.find("。")
                        if end_idx != -1:
                            run = new_para.add_run("    " + original_text[:end_idx + 1])
                            apply_format(run, 'arabic_number')
                            if end_idx + 1 < len(original_text):
                                run = new_para.add_run(original_text[end_idx + 1:])
                                apply_format(run, 'default')
                            else:
                                run = new_para.add_run("    " + original_text)
                                apply_format(run, 'default')
                            break
            # 检查是否包含"一是"到"十是"的句子
            elif any(f"{num}是" in original_text for num in ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]):
                # 找到"X是"的位置和句号的位置
                for num in ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]:
                    pattern = f"{num}是"
                    if pattern in original_text:
                        start_idx = original_text.find(pattern)
                        end_idx = original_text.find("。", start_idx)
                        if end_idx != -1:  # 如果找到句号
                            run = new_para.add_run("      " + original_text[:start_idx])
                            # 添加前面的文本（如果有）
                            if start_idx > 0:
                                run.font.name = rules['content']['default']['font']
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), rules['content']['default']['font'])
                                run.font.bold = rules['content']['default']['bold']
                                run.font.size = Pt(rules['content']['default']['size'])
                            
                            # 添加"X是"到句号的部分（包含句号）
                            run = new_para.add_run(original_text[start_idx:end_idx + 1])
                            run.font.name = rules['content']['number_is']['font']
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), rules['content']['number_is']['font'])
                            run.font.bold = rules['content']['number_is']['bold']
                            run.font.size = Pt(rules['content']['default']['size'])
                            
                            # 添加剩余的文本（如果有）
                            if end_idx + 1 < len(original_text):
                                run = new_para.add_run(original_text[end_idx + 1:])
                                run.font.name = rules['content']['default']['font']
                                run._element.rPr.rFonts.set(qn('w:eastAsia'), rules['content']['default']['font'])
                                run.font.bold = rules['content']['default']['bold']
                                run.font.size = Pt(rules['content']['default']['size'])
                            break
                        else:  # 如果没找到句号，按原样处理整个段落
                            run = new_para.add_run("    " + original_text)
                            run.font.name = rules['content']['default']['font']
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), rules['content']['default']['font'])
                            run.font.bold = rules['content']['default']['bold']
                            run.font.size = Pt(rules['content']['default']['size'])
            # 检查是否以中文数字加顿号开头
            elif any(original_text.startswith(f"{num}、") for num in ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]):
                run = new_para.add_run("    " + original_text)
                apply_format(run, 'number_dot')
            else:
                run = new_para.add_run("    " + original_text)
                apply_format(run, 'default')
            
            run.font.size = Pt(rules['content']['default']['size'])  # 使用默认字号
    
    # 保存修改后的文档
    doc.save(doc_path)


